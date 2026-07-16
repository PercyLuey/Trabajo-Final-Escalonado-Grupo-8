from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]; OUT=ROOT/'reports'/'Informe_T2_Grupo_8.docx'; FIG=ROOT/'outputs'/'figures'
BLUE='2E74B5'; DARK='1F4D78'; LIGHT='E8EEF5'; GRAY='666666'

def font(run,size=11,bold=False,color='000000',italic=False):
    run.font.name='Calibri'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)

def shade(cell,fill):
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(shd)

def set_cell_width(cell,dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW')) or OxmlElement('w:tcW'); tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa');
    if tcW.getparent() is None: tcPr.append(tcW)

def table(rows,widths):
    t=doc.add_table(rows=1, cols=len(widths)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False; t.style='Table Grid'
    tblPr=t._tbl.tblPr; tblW=tblPr.find(qn('w:tblW')); tblW.set(qn('w:w'),'9360'); tblW.set(qn('w:type'),'dxa')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    grid=t._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths: gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); grid.append(gc)
    for j,v in enumerate(rows[0]):
        set_cell_width(t.rows[0].cells[j],widths[j]); t.rows[0].cells[j].text=v; shade(t.rows[0].cells[j],LIGHT)
        for r in t.rows[0].cells[j].paragraphs[0].runs: font(r,9,bold=True,color=DARK)
    for row in rows[1:]:
        cells=t.add_row().cells
        for j,v in enumerate(row):
            set_cell_width(cells[j],widths[j]); cells[j].text=str(v); cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for r in cells[j].paragraphs[0].runs: font(r,9)
    doc.add_paragraph().paragraph_format.space_after=Pt(1); return t

def p(text='',bold=False,italic=False,color='000000',align=None,after=8):
    x=doc.add_paragraph(); x.paragraph_format.space_after=Pt(after); x.paragraph_format.line_spacing=1.333
    if align is not None: x.alignment=align
    font(x.add_run(text),11,bold,color,italic); return x

def bullet(text):
    x=doc.add_paragraph(style='List Bullet'); x.paragraph_format.left_indent=Inches(.375); x.paragraph_format.first_line_indent=Inches(-.194); x.paragraph_format.space_after=Pt(4); x.paragraph_format.line_spacing=1.208; font(x.add_run(text)); return x

def heading(text,level=1): doc.add_heading(text,level=level)
def figure(name,caption,width=6.25):
    doc.add_picture(str(FIG/name),width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c=p(caption,italic=True,color=GRAY,align=WD_ALIGN_PARAGRAPH.CENTER,after=10); c.runs[0].font.size=Pt(9)

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.492)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.333
for n,size,before,after,color in [('Title',30,0,8,DARK),('Heading 1',16,18,10,BLUE),('Heading 2',13,12,6,BLUE),('Heading 3',12,8,4,DARK)]:
    s=styles[n]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=(n!='Title'); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
header=sec.header.paragraphs[0]; header.text='Aplicaciones de IA en Estructuras | Grupo 8'; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: font(r,9,color=GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

p('TRABAJO FINAL ESCALONADO · ENTREGA T2',bold=True,color=BLUE,align=WD_ALIGN_PARAGRAPH.CENTER,after=20)
t=doc.add_paragraph(style='Title'); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(t.add_run('Predicción automática de respuesta sísmica mediante Machine Learning'),30,bold=False,color=DARK)
p('Análisis exploratorio de datos y plan algorítmico',italic=True,color=GRAY,align=WD_ALIGN_PARAGRAPH.CENTER,after=34)
p('Curso: Aplicaciones de IA en Estructuras',bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
p('Docente: Ing. Kurt Soncco Sinchi',align=WD_ALIGN_PARAGRAPH.CENTER)
p('Julio de 2026',align=WD_ALIGN_PARAGRAPH.CENTER,after=28)
table([['Integrante','Participación declarada en T1'],['Percy Luey Espinoza','Aportó al desarrollo'],['Delarc Dario Ayala Cacha','No aportó al desarrollo'],['Martin Eduardo Suarez Cruz','Aportó al desarrollo']],[4200,5160])
p('Repositorio: Trabajo-Final-Escalonado-Grupo-8',italic=True,color=GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

heading('Resumen ejecutivo')
p('Este informe consolida el planteamiento T1 y desarrolla la entrega T2: incorporación del dataset completo, control de calidad, análisis exploratorio reproducible en Python, plan de algoritmos y criterios de evaluación. El problema se formula como un metamodelo de preevaluación de respuesta sísmica de edificaciones, orientado a priorizar escenarios de análisis en ingeniería civil.')
p('Hallazgo principal.',bold=True,color=DARK,after=3)
p('El archivo completo contiene 1,000 observaciones y 26 variables, sin nulos ni duplicados. Sin embargo, las correlaciones de Spearman entre las entradas y cada respuesta son extremadamente pequeñas (máximo absoluto aproximado de 0.066). Además, la probabilidad de colapso es exactamente cien veces el índice de daño. Estos resultados sugieren datos sintéticos con respuestas débilmente vinculadas a los predictores y una redundancia determinística entre dos objetivos. Por ello, T3 debe demostrar valor frente a una línea base y no asumir de antemano que existe señal predictiva útil.')
heading('1. Problema de ingeniería estructural')
heading('1.1 Contexto y objetivo',2)
p('La evaluación sísmica rigurosa requiere modelos estructurales, propiedades de materiales, acciones, condiciones de sitio y procedimientos compatibles con la normativa aplicable. Para estudios preliminares de múltiples alternativas, un metamodelo puede aproximar salidas de simulaciones y reducir el número de análisis detallados. El objetivo es estimar, con incertidumbre explícita, deriva máxima de entrepiso y desplazamiento máximo de techo; otras respuestas se tratarán como objetivos secundarios.')
heading('1.2 Alcance responsable',2)
bullet('Uso previsto: clasificación preliminar de escenarios y apoyo académico a la comparación de modelos de ML.')
bullet('Fuera de alcance: diseño, certificación, evaluación de seguridad o reemplazo de análisis modal, tiempo-historia o no lineal.')
bullet('Condición para uso futuro: validación externa con modelos estructurales trazables o mediciones reales y revisión de un ingeniero competente.')

heading('2. Dataset y trazabilidad')
p('Fuente: “Pre-Earthquake Prediction Dataset”, publicado por Ziya en Kaggle bajo licencia CC0. El archivo descargado es seismic_data.csv y ocupa aproximadamente 415 kB. La ficha declara que integra amenaza, movimiento del suelo, sitio, atributos estructurales e indicadores de respuesta. URL: https://www.kaggle.com/datasets/ziya07/pre-earthquake-prediction-dataset')
table([['Grupo','Variables principales'],['Amenaza y movimiento','Zona sísmica, PGA, PGV, PGD, aceleración espectral, magnitud, distancia a falla y frecuencia'],['Sitio','Tipo de suelo y factor de amplificación'],['Estructura','Altura, pisos, material, cimentación, frecuencia natural, amortiguamiento, masa, rigideces y sistema lateral'],['Objetivos','Deriva, desplazamiento de techo, cortante basal, aceleración, daño y colapso']],[2200,7160])
p('La versión parcial existente en el repositorio contenía solo diez variables de amenaza y sitio. Para T2 se incorporó el archivo completo de 26 columnas en data/raw, conservando el archivo parcial como antecedente.',italic=True,color=GRAY)

heading('3. Metodología del EDA')
p('El script src/eda.py implementa un flujo reproducible: lectura del CSV, clasificación de tipos, conteo de nulos y duplicados, resúmenes descriptivos, inspección de categorías, correlación de Spearman y visualizaciones. Spearman se seleccionó porque detecta relaciones monótonas sin exigir normalidad y es menos sensible a escalas diferentes.')
table([['Control','Resultado','Decisión'],['Dimensiones','1,000 × 26','Trabajar con el archivo completo'],['Tipos','21 numéricas; 5 categóricas','Codificación dentro del pipeline'],['Valores nulos','0','Mantener imputador preventivo para datos futuros'],['Duplicados','0','No eliminar registros'],['Memoria','0.427 MB','Ejecución local/Colab viable']],[1900,1800,5660])

heading('4. Resultados exploratorios')
heading('4.1 Distribución de respuestas',2)
table([['Objetivo','Media','Desv. estándar','Mín.–Máx.'],['Deriva máxima (%)','2.514','1.428','0.104–4.994'],['Desplazamiento de techo (m)','1.011','0.575','0.010–1.999'],['Cortante basal (kN)','2,523.310','1,432.974','51.648–4,997.439'],['Aceleración estructural (m/s²)','2.499','1.432','0.105–4.995'],['Índice de daño','0.512','0.286','0.001–1.000'],['Probabilidad de colapso (%)','51.177','28.553','0.123–99.952']],[3600,1500,1900,2360])
figure('01_distribucion_objetivos.png','Figura 1. Distribuciones de las seis respuestas provistas por el dataset.')
heading('4.2 Correlaciones y coherencia física',2)
p('La mayor correlación absoluta observada entre una entrada y un objetivo fue aproximadamente 0.066 (rigidez axial frente a índice de daño/colapso). Para deriva, la mayor fue PGA con ρ≈−0.055; para desplazamiento de techo, PGA con ρ≈0.056. Estos valores y algunos signos poco intuitivos no prueban ausencia total de relaciones no lineales, pero sí advierten que los modelos podrían no superar una predicción constante fuera de muestra.')
figure('02_matriz_correlacion.png','Figura 2. Matriz de correlaciones de Spearman. La ausencia de bandas intensas entre entradas y objetivos indica señal monótona débil.',5.8)
p('Se comprobó además que Collapse Probability (%) = 100 × Damage Index para todas las filas, con discrepancia numérica máxima del orden de 10⁻¹⁴. Ambos objetivos no deben usarse simultáneamente como predictores y salidas; bastará modelar uno o reportar la relación determinística.')
figure('03_respuesta_por_sistema.png','Figura 3. Comparación exploratoria de respuesta por material y sistema resistente.')
figure('04_pga_aceleracion.png','Figura 4. PGA frente a aceleración estructural, diferenciada por tipo de suelo.',5.8)

heading('5. Plan algorítmico para T3')
table([['Modelo','Función en el estudio','Justificación'],['DummyRegressor','Línea base obligatoria','Determina si existe ganancia real frente a predecir la media/mediana'],['Ridge','Base lineal regularizada','Modelo interpretable y estable ante colinealidad'],['Random Forest','Modelo no lineal de ensamble','Captura interacciones y no requiere escalamiento'],['Gradient Boosting','Alternativa no lineal','Puede representar relaciones complejas con regularización'],['MLP','Extensión condicionada','Solo si la validación muestra señal suficiente; 1,000 filas elevan el riesgo de sobreajuste']],[1900,2600,4860])
heading('5.1 Pipeline y control de fuga',2)
bullet('Excluir las seis salidas “Predicted” del conjunto X; en particular, impedir que daño prediga colapso o viceversa.')
bullet('Separar 80% entrenamiento y 20% prueba antes de ajustar transformaciones; usar validación cruzada de cinco pliegues solo en entrenamiento.')
bullet('Imputar mediana/moda, One-Hot Encoding para categorías y escalamiento para Ridge/MLP mediante ColumnTransformer y Pipeline.')
bullet('Repetir el proceso con varias semillas y conservar la prueba final intacta hasta seleccionar el modelo.')
heading('5.2 Métricas y resultados previstos',2)
p('MAE será la métrica principal por su interpretación en unidades de ingeniería; RMSE resaltará errores grandes y R² será complementaria. MAPE se usará únicamente si no hay valores cercanos a cero. No se prometen cifras antes de entrenar. El resultado esperado más responsable es uno de dos: (a) un ensamble supera de forma estable a las líneas base y permite un metamodelo exploratorio; o (b) ningún modelo generaliza, lo que confirmaría que el dataset no contiene señal suficiente para las salidas. Ambos resultados son científicamente válidos.')

heading('6. Marco PCS: predictibilidad, computabilidad y estabilidad')
table([['Pilar','Aplicación en T2/T3'],['Predictibilidad','Prueba no vista, comparación con DummyRegressor y métricas con intervalos/dispersión entre pliegues.'],['Computabilidad','CSV pequeño, código Python versionado, dependencias declaradas y salidas regenerables.'],['Estabilidad','Variar semillas, codificaciones y subconjuntos; comparar métricas e importancia por permutación.']],[1900,7460])
p('La revisión PCS también exige cuestionar la procedencia de las etiquetas. La palabra “Predicted” indica que las respuestas no son observaciones de campo; por tanto, el modelo aprendería a emular un generador previo. Esta inferencia debe validarse con el autor o documentación adicional antes de formular conclusiones físicas.')

heading('7. Conclusiones del avance T2')
bullet('Se actualizó el repositorio con el dataset completo y el informe T1 original.')
bullet('El EDA es ejecutable y produce tablas y cuatro figuras sin intervención manual.')
bullet('La calidad formal es alta en nulos y duplicados, pero la utilidad predictiva es incierta por correlaciones débiles y posible naturaleza sintética.')
bullet('Se definió un plan algorítmico con líneas base, control de fuga, validación cruzada, métricas físicas y análisis de estabilidad.')
bullet('El proyecto se mantiene como metamodelo académico; cualquier aplicación profesional requiere evidencia externa y verificación normativa.')

heading('8. Bitácora de uso de inteligencia artificial')
p('Herramienta: Codex, modelo de OpenAI, utilizado como asistente para estructurar el repositorio, generar y depurar código Python, analizar la calidad del dataset, redactar el informe y preparar el speech. Prompt inicial resumido: avanzar el repositorio del Grupo 8 hasta T2 usando el dataset indicado, incorporar datos e informe T1, generar EDA, plan de algoritmos, resultados previstos, Word y speech desde la perspectiva de ingeniería civil.')
p('Iteraciones registradas: 4 refinamientos principales: (1) inspección de rúbrica/T1; (2) sustitución de la versión parcial por el archivo completo; (3) corrección del entorno y ejecución del EDA; (4) revisión técnica de correlaciones, redundancia de objetivos y alcance profesional. Validación humana requerida: el equipo debe ejecutar el script, revisar cifras, confirmar la procedencia de las etiquetas y asumir responsabilidad por la entrega.')

heading('Referencias')
p('Kaggle. (s. f.). Pre-Earthquake Prediction Dataset. https://www.kaggle.com/datasets/ziya07/pre-earthquake-prediction-dataset')
p('Yu, B., & Barter, R. L. (2024). Veridical Data Science: The Practice of Responsible Data Analysis and Decision Making. MIT Press. https://vdsbook.com/')
p('Breiman, L. (2001). Random Forests. Machine Learning, 45, 5–32. https://doi.org/10.1023/A:1010933404324')
p('Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. The Annals of Statistics, 29(5), 1189–1232. https://doi.org/10.1214/aos/1013203451')
p('OpenAI. (2026). Codex [Large language model]. https://openai.com/codex/')

heading('Anexo A. Reproducibilidad y evidencias')
table([['Elemento','Ubicación','Evidencia'],['Datos originales','data/raw/seismic_data.csv','1,000 filas y 26 columnas'],['Código EDA','src/eda.py','Genera controles, CSV de resumen y cuatro figuras'],['Salidas tabulares','outputs/*.csv y eda_summary.json','Calidad, descriptivos y correlaciones'],['Visualizaciones','outputs/figures/*.png','Distribuciones, correlaciones y comparaciones'],['Plan y exposición','reports/*.md','Plan algorítmico y speech de cinco minutos']],[1900,3400,4060])
p('Comando mínimo de reproducción: instalar requirements.txt y ejecutar python src/eda.py desde la raíz del repositorio. La ejecución debe terminar sin errores y regenerar los archivos de outputs. Antes de T3, el equipo deberá añadir pruebas automáticas para verificar dimensiones, ausencia de fuga entre X e y y relación determinística entre daño y colapso.')

doc.save(OUT); print(OUT)
